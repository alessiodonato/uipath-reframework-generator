#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════════╗
║         UiPath ReFramework Generator — Powered by Claude AI     ║
║  Generates a complete, ready-to-open UiPath Studio project       ║
║  from a Process Definition Document (PDF or DOCX).               ║
╚══════════════════════════════════════════════════════════════════╝

Usage:
    # Interactive mode (recommended for first use)
    python generate_reframework.py

    # Direct mode
    python generate_reframework.py --input MyProcess_PDD.pdf --output ./output/

    # Preview extracted metadata without generating files
    python generate_reframework.py --input MyProcess_PDD.pdf --metadata-only

    # Use a specific API key
    python generate_reframework.py --input MyProcess_PDD.pdf --api-key sk-ant-...

Requirements:
    Set ANTHROPIC_API_KEY environment variable, or pass --api-key flag.
    Dependencies are auto-installed on first run if missing.
"""

import argparse
import json
import os
import re
import subprocess
import sys
import zipfile
from datetime import datetime
from pathlib import Path


# ─── CONSTANTS ───────────────────────────────────────────────────────────────

VERSION = "2.0.0"
TEMPLATES_DIR = Path(__file__).parent.parent / "assets" / "xaml-templates"
SEQUENCE_TEMPLATES_DIR = Path(__file__).parent.parent / "assets" / "sequence-template"
CLAUDE_MODEL = "claude-sonnet-4-20250514"

# Project variants
VARIANTS = {
    "reframework": "ReFramework with State Machine (default)",
    "dispatcher": "ReFramework Dispatcher (loads data into queue)",
    "performer": "ReFramework Performer (processes queue items)",
    "sequence": "Simple Sequence (linear, no transaction loop)",
}

REQUIRED_PACKAGES = {
    "pdfminer": "pdfminer.six",
    "docx": "python-docx",
    "openpyxl": "openpyxl",
    "anthropic": "anthropic",
}


# ─── STEP 0: AUTO-INSTALL DEPENDENCIES ──────────────────────────────────────

def ensure_dependencies():
    """Auto-install missing Python packages on first run."""
    missing = []
    for import_name, pip_name in REQUIRED_PACKAGES.items():
        try:
            __import__(import_name)
        except ImportError:
            missing.append(pip_name)

    if not missing:
        return

    print(f"📦 Installing missing dependencies: {', '.join(missing)}")
    try:
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "--quiet"] + missing,
            stderr=subprocess.DEVNULL
        )
        print("  ✓ Dependencies installed\n")
    except subprocess.CalledProcessError:
        # Fallback: try with --break-system-packages (Linux/Claude environment)
        try:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", "--quiet",
                 "--break-system-packages"] + missing,
                stderr=subprocess.DEVNULL
            )
            print("  ✓ Dependencies installed\n")
        except subprocess.CalledProcessError as e:
            print(f"  ✗ Could not auto-install. Please run manually:")
            print(f"    pip install {' '.join(missing)}")
            sys.exit(1)


# ─── STEP 1: EXTRACT TEXT FROM DOCUMENT ─────────────────────────────────────

def extract_text_from_pdf(path: str) -> str:
    from pdfminer.high_level import extract_text
    text = extract_text(path)
    if not text or not text.strip():
        raise ValueError("PDF appears to be empty or scanned (no extractable text). "
                         "Please use a text-based PDF.")
    return text


def extract_text_from_docx(path: str) -> str:
    import docx
    doc = docx.Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    text = "\n".join(parts)
    if not text.strip():
        raise ValueError("DOCX appears to be empty.")
    return text


def extract_document_text(input_path: str) -> str:
    ext = Path(input_path).suffix.lower()
    if not Path(input_path).exists():
        raise FileNotFoundError(f"File not found: {input_path}")
    if ext == ".pdf":
        return extract_text_from_pdf(input_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(input_path)
    else:
        raise ValueError(f"Unsupported format: {ext}. Supported: PDF, DOCX")


# ─── STEP 2: CALL CLAUDE API — METADATA EXTRACTION ──────────────────────────

EXTRACTION_SYSTEM_PROMPT = """You are a senior UiPath RPA architect with 10+ years of experience.
You will receive the text of a Process Definition Document (PDD) or any process analysis document.

Your task is to extract ALL information needed to scaffold a production-ready UiPath ReFramework project.

Return ONLY a valid JSON object — no markdown fences, no preamble, no explanation.

JSON structure:
{
  "process_name": "<PascalCase, no spaces, e.g. InvoiceProcessing>",
  "process_description": "<clear one-sentence description of what the robot does>",
  "applications": ["<AppName>"],
  "transaction_source": "<Queue | Excel | Database | API>",
  "queue_name": "<OrchestratorQueueName or empty string if not Queue>",
  "transaction_item_fields": [
    { "name": "<fieldName>", "type": "<String|Integer|Boolean|DateTime>", "description": "<what this field contains>" }
  ],
  "process_steps": [
    {
      "id": "Step01",
      "name": "<PascalCase step name, e.g. LoginToSAP>",
      "description": "<precise description of what this step does>",
      "app": "<application name>",
      "pseudo_steps": [
        "<numbered sub-step, e.g. 1. Navigate to URL from Config('SAP_URL')>",
        "<2. Enter username from credential asset>",
        "<3. Click Login button>",
        "<4. Verify dashboard is visible, else throw ApplicationException>"
      ],
      "config_keys_used": ["<list of Config keys read in this step, e.g. 'SAP_URL', 'SAP_CredentialAsset'>"],
      "output_variables": [
        { "name": "<varName>", "type": "<String|Integer|Boolean|DataTable>", "description": "<what it contains>" }
      ],
      "throws": [
        { "exception_type": "<BusinessRuleException | ApplicationException>", "condition": "<when thrown>" }
      ],
      "business_rule": "<any business rule or validation this step must enforce, or empty string>"
    }
  ],
  "business_exceptions": [
    {
      "name": "<ExceptionName e.g. InvoiceAlreadyProcessedException>",
      "condition": "<precise condition when this should be thrown>",
      "suggested_step": "<which process step should throw this>"
    }
  ],
  "system_exceptions": [
    {
      "name": "<ExceptionName e.g. SAPLoginFailedException>",
      "condition": "<when this system exception could occur>",
      "recovery_hint": "<suggested recovery action>"
    }
  ],
  "config_settings": [
    { "name": "<SettingName>", "value": "<default or placeholder value>", "type": "<Setting|Constant|Asset>", "description": "<what it's used for>" }
  ],
  "output_report": {
    "enabled": true,
    "fields": ["<field from transaction_item_fields to report on>"]
  },
  "max_retry_number": 3,
  "process_type": "<Transactional | Linear>"
}

Rules:
- process_name: PascalCase, no spaces, no special characters
- Default transaction_source to "Queue" if not specified
- Default max_retry_number to 3 unless document explicitly states otherwise
- For each application: include URL setting and credential asset setting in config_settings
- process_steps must be ONLY business steps — never include Init, GetTransactionData, SetTransactionStatus
- pseudo_steps must be concrete and actionable — reference Config keys, credential assets, specific UI elements mentioned in the document
- config_keys_used: list ALL Config keys referenced in this step's pseudo_steps
- throws: list ALL exceptions this step might throw (BusinessRuleException for validation failures, ApplicationException for system failures)
- If the document mentions specific fields, forms, buttons or screens — include them in pseudo_steps
- business_exceptions: use UiPath naming convention (verb + noun + Exception)
- output_variables in each step = variables that downstream steps will need
- process_type = "Linear" only if the document explicitly describes a one-shot process with no transaction loop"""


def extract_metadata_via_api(document_text: str, api_key: str) -> dict:
    """Calls Claude API to extract structured ReFramework metadata from PDD text."""
    import anthropic

    client = anthropic.Anthropic(api_key=api_key)

    # Truncate text to avoid token limits while keeping most important content
    text_chunk = document_text[:20000]
    if len(document_text) > 20000:
        # Also include the end of the document (often contains exception tables)
        text_chunk += "\n\n[...]\n\n" + document_text[-3000:]

    message = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4096,
        system=EXTRACTION_SYSTEM_PROMPT,
        messages=[
            {
                "role": "user",
                "content": f"Here is the process document to analyze:\n\n---\n{text_chunk}\n---\n\nExtract the complete ReFramework metadata JSON."
            }
        ]
    )

    raw = message.content[0].text.strip()
    # Strip possible markdown fences (defensive)
    raw = re.sub(r"^```json\s*", "", raw, flags=re.MULTILINE)
    raw = re.sub(r"\s*```\s*$", "", raw, flags=re.MULTILINE)
    raw = raw.strip()

    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        raise ValueError(f"Claude returned invalid JSON: {e}\nRaw output:\n{raw[:500]}")


# ─── STEP 3: GENERATE XAML COMPONENTS ────────────────────────────────────────

def load_template(filename: str) -> str:
    path = TEMPLATES_DIR / filename
    if not path.exists():
        raise FileNotFoundError(f"Template not found: {path}\n"
                                f"Make sure you're running from the skill root directory.")
    return path.read_text(encoding="utf-8")


def apply_common_replacements(content: str, metadata: dict) -> str:
    today = datetime.now().strftime("%Y-%m-%d")
    apps_list = ", ".join(metadata.get("applications", []))
    return (content
        .replace("{{PROCESS_NAME}}", metadata["process_name"])
        .replace("{{PROCESS_DESCRIPTION}}", metadata["process_description"])
        .replace("{{MAX_RETRY_NUMBER}}", str(metadata.get("max_retry_number", 3)))
        .replace("{{QUEUE_NAME}}", metadata.get("queue_name", ""))
        .replace("{{TRANSACTION_SOURCE}}", metadata.get("transaction_source", "Queue"))
        .replace("{{APPLICATIONS_LIST}}", apps_list)
        .replace("{{CURRENT_DATE}}", today)
        .replace("{{VERSION}}", VERSION)
    )


def generate_process_step_invocation(step: dict, process_name: str) -> str:
    """Generates an InvokeWorkflowFile block for Process.xaml."""
    pseudo = " | ".join(step.get("pseudo_steps", [])[:3])  # first 3 in annotation
    annotation = f"{step['description']}"
    if pseudo:
        annotation += f"&#xA;Steps: {pseudo}"

    # Determine output arguments if step declares output variables
    out_args = ""
    for var in step.get("output_variables", []):
        vb_type = _vb_type(var["type"])
        out_args += f"""
        <OutArgument x:TypeArguments="{vb_type}" x:Key="out_{var['name']}">[{var['name']}]</OutArgument>"""

    return f"""
    <!-- ── {step['id']}: {step['name']} ── -->
    <ui:InvokeWorkflowFile
      DisplayName="{step['id']} - {step['name']}"
      sap2010:Annotation.AnnotationText="{annotation}"
      WorkflowFileName="Business\\{step['name']}.xaml"
      UnSafe="False">
      <ui:InvokeWorkflowFile.Arguments>
        <InArgument x:TypeArguments="ui:QueueItem" x:Key="in_TransactionItem">[in_TransactionItem]</InArgument>
        <InArgument x:TypeArguments="scg:Dictionary(x:String, x:Object)" x:Key="in_Config">[in_Config]</InArgument>{out_args}
      </ui:InvokeWorkflowFile.Arguments>
    </ui:InvokeWorkflowFile>"""


def generate_app_init_invocation(app: str) -> str:
    """Generates an InvokeWorkflowFile block for InitAllApplications.xaml."""
    return f"""
    <!-- ── Open {app} ── -->
    <ui:InvokeWorkflowFile
      DisplayName="Open {app}"
      sap2010:Annotation.AnnotationText="Initialize and open {app}.&#xA;URL: Config(&quot;{app}_URL&quot;) | Credential: Config(&quot;{app}_CredentialAsset&quot;)"
      WorkflowFileName="Business\\Open{app}.xaml"
      UnSafe="False">
      <ui:InvokeWorkflowFile.Arguments>
        <InArgument x:TypeArguments="scg:Dictionary(x:String, x:Object)" x:Key="in_Config">[in_Config]</InArgument>
      </ui:InvokeWorkflowFile.Arguments>
    </ui:InvokeWorkflowFile>"""


def generate_get_transaction_body(metadata: dict, process_name: str) -> str:
    """Generate the body for GetTransactionData.xaml with standardized log messages."""
    source = metadata.get("transaction_source", "Queue")
    queue = metadata.get("queue_name", "INPUT_QUEUE_NAME")

    if source == "Queue":
        return f"""
    <!-- ── Queue-based retrieval ── -->
    <ui:LogMessage Level="Trace"
      Message="['[{process_name}] GetTx - Fetching next transaction (attempt ' + in_TransactionNumber.ToString() + ')']"
      DisplayName="Log: GetTx - Fetching next transaction" />

    <ui:GetQueueItem
      DisplayName="Get Queue Item from '{queue}'"
      sap2010:Annotation.AnnotationText="Retrieves the next pending item from Orchestrator queue.&#xA;Returns Nothing when queue is empty → triggers End Process.&#xA;&#xA;Reads: Config('OrchestratorQueueName') | Output: out_TransactionItem (QueueItem) | Throws: N/A"
      QueueName="[If(String.IsNullOrEmpty(in_Config(&quot;OrchestratorQueueName&quot;).ToString()), &quot;{queue}&quot;, in_Config(&quot;OrchestratorQueueName&quot;).ToString())]"
      TimeoutMS="30000"
      Result="[out_TransactionItem]" />

    <If Condition="[out_TransactionItem IsNot Nothing]" DisplayName="Check if transaction retrieved">
      <If.Then>
        <ui:LogMessage Level="Info"
          Message="['[{process_name}] GetTx - Transaction retrieved: ' + out_TransactionItem.Reference]"
          DisplayName="Log: GetTx - Transaction retrieved" />
      </If.Then>
      <If.Else>
        <ui:LogMessage Level="Info"
          Message="'[{process_name}] GetTx - No more transactions'"
          DisplayName="Log: GetTx - No more transactions" />
      </If.Else>
    </If>"""
    elif source == "Excel":
        return f"""
    <!-- ── Excel/DataTable-based retrieval ── -->
    <ui:LogMessage Level="Trace"
      Message="['[{process_name}] GetTx - Fetching next transaction (attempt ' + in_TransactionNumber.ToString() + ')']"
      DisplayName="Log: GetTx - Fetching next transaction" />

    <!-- TODO: Load DataTable from Excel in InitAllApplications, pass as variable -->
    <If Condition="[in_TransactionNumber &lt;= TransactionData.Rows.Count]"
      DisplayName="Check if more rows exist"
      sap2010:Annotation.AnnotationText="Reads: TransactionData (DataTable) | Output: out_TransactionItem | Throws: N/A">
      <If.Then>
        <Sequence DisplayName="Get DataRow as TransactionItem">
          <Assign DisplayName="Get row">
            <Assign.To><OutArgument x:TypeArguments="ui:QueueItem">[out_TransactionItem]</OutArgument></Assign.To>
            <!-- TODO: Map DataRow fields to a QueueItem or use DataRow directly -->
            <!-- Row index: in_TransactionNumber - 1 -->
            <!-- Access fields: TransactionData.Rows(in_TransactionNumber - 1)("FieldName").ToString() -->
          </Assign>
          <ui:LogMessage Level="Info"
            Message="['[{process_name}] GetTx - Transaction retrieved: Row ' + in_TransactionNumber.ToString()]"
            DisplayName="Log: GetTx - Transaction retrieved" />
        </Sequence>
      </If.Then>
      <If.Else>
        <Sequence DisplayName="No more rows — signal End">
          <Assign DisplayName="Set to Nothing">
            <Assign.To><OutArgument x:TypeArguments="ui:QueueItem">[out_TransactionItem]</OutArgument></Assign.To>
            <Assign.Value><InArgument x:TypeArguments="ui:QueueItem">Nothing</InArgument></Assign.Value>
          </Assign>
          <ui:LogMessage Level="Info"
            Message="'[{process_name}] GetTx - No more transactions'"
            DisplayName="Log: GetTx - No more transactions" />
        </Sequence>
      </If.Else>
    </If>"""
    else:
        return f"""
    <!-- ── {source}-based retrieval ── -->
    <ui:LogMessage Level="Trace"
      Message="['[{process_name}] GetTx - Fetching next transaction (attempt ' + in_TransactionNumber.ToString() + ')']"
      DisplayName="Log: GetTx - Fetching next transaction" />

    <!-- TODO: Implement {source} data source retrieval -->
    <!-- Pattern: set out_TransactionItem = Nothing when no more data available -->
    <!--
      For Database: use Execute Query activity, iterate rows using in_TransactionNumber as offset
      For API: use HTTP Request activity, paginate using in_TransactionNumber as page index
    -->
    <Sequence DisplayName="TODO: Get transaction from {source}"
      sap2010:Annotation.AnnotationText="TODO: implement - Retrieve transaction from {source}&#xA;&#xA;Reads: N/A | Output: out_TransactionItem | Throws: ApplicationException if retrieval fails">
      <Assign DisplayName="TODO: Get transaction from {source}">
        <Assign.To><OutArgument x:TypeArguments="ui:QueueItem">[out_TransactionItem]</OutArgument></Assign.To>
        <Assign.Value><InArgument x:TypeArguments="ui:QueueItem">Nothing</InArgument></Assign.Value>
      </Assign>
      <ui:LogMessage Level="Info"
        Message="'[{process_name}] GetTx - No more transactions'"
        DisplayName="Log: GetTx - No more transactions" />
    </Sequence>"""


def _vb_type(type_str: str) -> str:
    """Map simple type names to VB.NET XAML type arguments (for x:TypeArguments)."""
    mapping = {
        "String": "x:String",
        "Integer": "x:Int32",
        "Int": "x:Int32",
        "Boolean": "x:Boolean",
        "Bool": "x:Boolean",
        "DateTime": "s:DateTime",
        "Date": "s:DateTime",
        "DataTable": "sd:DataTable",
        "Table": "sd:DataTable",
        "Double": "x:Double",
        "Decimal": "x:Double",
        "Float": "x:Double",
    }
    return mapping.get(type_str, mapping.get(type_str.capitalize(), "x:String"))


def map_type(raw_type: str) -> tuple[str, str]:
    """
    Map raw type string from extraction to XAML x:Type declaration.
    Returns (xaml_type, todo_comment) where todo_comment is empty if type is recognized.
    """
    if not raw_type:
        return "x:Type p:String", "<!-- TODO: verify type (was empty) -->"

    normalized = raw_type.lower().strip()

    # Direct mappings
    type_map = {
        "string": ("x:Type p:String", ""),
        "str": ("x:Type p:String", ""),
        "int": ("x:Type p:Int32", ""),
        "integer": ("x:Type p:Int32", ""),
        "bool": ("x:Type p:Boolean", ""),
        "boolean": ("x:Type p:Boolean", ""),
        "datatable": ("x:Type sd:DataTable", ""),
        "table": ("x:Type sd:DataTable", ""),
        "data table": ("x:Type sd:DataTable", ""),
        "queueitem": ("x:Type uia:QueueItem", ""),
        "queue item": ("x:Type uia:QueueItem", ""),
        "dict": ("s:Dictionary(x:Type p:String, x:Type p:Object)", ""),
        "dictionary": ("s:Dictionary(x:Type p:String, x:Type p:Object)", ""),
        "config": ("s:Dictionary(x:Type p:String, x:Type p:Object)", ""),
        "datetime": ("x:Type s:DateTime", ""),
        "date": ("x:Type s:DateTime", ""),
        "double": ("x:Type p:Double", ""),
        "decimal": ("x:Type p:Double", ""),
        "float": ("x:Type p:Double", ""),
    }

    if normalized in type_map:
        return type_map[normalized]

    # Check for list/array/collection patterns
    if any(kw in normalized for kw in ["list", "array", "collection"]):
        if "string" in normalized:
            return "s:IEnumerable(x:Type p:String)", ""
        return "s:IEnumerable(x:Type p:String)", f"<!-- TODO: verify collection element type for '{raw_type}' -->"

    # Fallback: String with TODO
    return "x:Type p:String", f"<!-- TODO: verify type for '{raw_type}' -->"


# ─── STEP 4: GENERATE BUSINESS WORKFLOW STUBS ────────────────────────────────

def _build_annotation(step: dict) -> str:
    """Build rich annotation string: Reads: ... | Output: ... | Throws: ..."""
    # Config keys used
    config_keys = step.get("config_keys_used", [])
    if config_keys:
        reads_part = "Reads: " + ", ".join(f"Config('{k}')" for k in config_keys)
    else:
        reads_part = "Reads: N/A"

    # Output variables
    out_vars = step.get("output_variables", [])
    if out_vars:
        output_part = "Output: " + ", ".join(f"{v['name']} ({v.get('type', 'String')})" for v in out_vars)
    else:
        output_part = "Output: N/A"

    # Throws
    throws = step.get("throws", [])
    if throws:
        throws_part = "Throws: " + ", ".join(
            f"{t['exception_type']} if {t['condition']}" for t in throws
        )
    else:
        throws_part = "Throws: N/A"

    return f"{reads_part} | {output_part} | {throws_part}"


def _escape_xml_attr(text: str) -> str:
    """Escape text for use in XML attributes."""
    return (text
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("\n", "&#xA;"))


def generate_business_step_xaml(step: dict, process_name: str) -> str:
    """
    Generates Business/StepName.xaml with:
    - Rich annotations (Reads/Output/Throws format)
    - Pseudo-steps as executable structure (LogMessage + placeholder Sequence)
    - Declared variables
    - Log bookends
    """
    # Build argument declarations with rich annotations
    annotation_summary = _build_annotation(step)

    in_args = f"""    <x:Property Name="in_TransactionItem" Type="InArgument(ui:QueueItem)"
      sap2010:Annotation.AnnotationText="Current transaction item. Access fields via in_TransactionItem.SpecificContent(&amp;quot;FieldName&amp;quot;).ToString()" />
    <x:Property Name="in_Config" Type="InArgument(scg:Dictionary(x:String, x:Object))"
      sap2010:Annotation.AnnotationText="Configuration dictionary. {_escape_xml_attr(annotation_summary)}" />"""

    out_arg_declarations = ""
    variables_xml = ""

    for var in step.get("output_variables", []):
        vb_type = _vb_type(var["type"])
        xaml_type, type_todo = map_type(var["type"])
        desc = _escape_xml_attr(var.get('description', ''))
        out_arg_declarations += f"""
    <x:Property Name="out_{var['name']}" Type="OutArgument({vb_type})"
      sap2010:Annotation.AnnotationText="{desc}" />{type_todo}"""
        variables_xml += f"""
      <Variable x:TypeArguments="{vb_type}" Name="{var['name']}" />"""

    # Build pseudo-steps as executable structure (LogMessage + Sequence placeholder with TODO annotation)
    pseudo_steps_xml = ""
    for idx, ps in enumerate(step.get("pseudo_steps", []), 1):
        ps_escaped = _escape_xml_attr(ps)
        pseudo_steps_xml += f"""
    <!-- ── Pseudo-step {idx} ── -->
    <ui:LogMessage Level="Trace"
      Message="'[{process_name}] {step['name']} - Step {idx}: {_escape_xml_attr(ps)}'"
      DisplayName="Log: Step {idx}" />
    <Sequence DisplayName="{ps_escaped}"
      sap2010:Annotation.AnnotationText="TODO: implement - {ps_escaped}">
      <!-- TODO: implement - {ps} -->
      <!-- Placeholder: replace this Sequence content with actual UiPath activities -->
    </Sequence>
"""

    if not pseudo_steps_xml:
        pseudo_steps_xml = """
    <!-- No pseudo-steps defined — add implementation activities here -->
    <Sequence DisplayName="TODO: Add implementation"
      sap2010:Annotation.AnnotationText="TODO: implement - No pseudo-steps were extracted from PDD">
      <!-- TODO: Add implementation activities -->
    </Sequence>
"""

    # Business rule as a validation block
    business_rule_xml = ""
    if step.get("business_rule"):
        rule_escaped = _escape_xml_attr(step['business_rule'])
        business_rule_xml = f"""
    <!-- ══════════════════════════════════════════════════════════════ -->
    <!-- BUSINESS RULE VALIDATION                                       -->
    <!-- ══════════════════════════════════════════════════════════════ -->
    <Sequence DisplayName="Validate: {rule_escaped}"
      sap2010:Annotation.AnnotationText="TODO: implement - Validate business rule: {rule_escaped}&#xA;Throw BusinessRuleException if rule is violated.">
      <!-- Rule: {step['business_rule']} -->
      <!-- TODO: Implement validation. Throw BusinessRuleException if rule is violated.
           Example:
           If <condition violates rule> Then
             Throw New BusinessRuleException("{step['business_rule']}")
           End If
      -->
    </Sequence>
"""

    return f"""<?xml version="1.0" encoding="utf-8"?>
<!--
  {step['name']}.xaml
  {step['description']}
  Application: {step['app']}
  Generated by ReFramework Generator v{VERSION} on {datetime.now().strftime("%Y-%m-%d")}
-->
<Activity
  mc:Ignorable="sap sap2010"
  x:Class="UiPath.ReFramework.Business.{step['name']}"
  xmlns="http://schemas.microsoft.com/netfx/2009/xaml/activities"
  xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
  xmlns:mva="clr-namespace:Microsoft.VisualBasic.Activities;assembly=System.Activities"
  xmlns:p="clr-namespace:System;assembly=mscorlib"
  xmlns:s="clr-namespace:System;assembly=mscorlib"
  xmlns:sap2010="http://schemas.microsoft.com/netfx/2010/xaml/activities/presentation"
  xmlns:scg="clr-namespace:System.Collections.Generic;assembly=mscorlib"
  xmlns:sd="clr-namespace:System.Data;assembly=System.Data"
  xmlns:ui="http://schemas.uipath.com/workflow/activities"
  xmlns:uia="clr-namespace:UiPath.Core;assembly=UiPath.Core"
  xmlns:x="http://schemas.microsoft.com/winfx/2006/xaml">
  <x:Members>
{in_args}{out_arg_declarations}
  </x:Members>
  <mva:VisualBasic.Settings>
    <mva:VisualBasicSettings>
      <mva:VisualBasicSettings.ImportedNamespaces>
        <mva:VisualBasicImportReference Assembly="mscorlib" Import="System" />
        <mva:VisualBasicImportReference Assembly="System.Data" Import="System.Data" />
        <mva:VisualBasicImportReference Assembly="UiPath.Core" Import="UiPath.Core" />
      </mva:VisualBasicSettings.ImportedNamespaces>
    </mva:VisualBasicSettings>
  </mva:VisualBasic.Settings>

  <Sequence
    sap2010:Annotation.AnnotationText="{_escape_xml_attr(step['description'])}&#xA;&#xA;{_escape_xml_attr(annotation_summary)}"
    DisplayName="{step['name']}">
    <Sequence.Variables>{variables_xml}
    </Sequence.Variables>

    <ui:LogMessage Level="Trace"
      Message="'[{process_name}] {step['name']} - Started | TxRef: ' + in_TransactionItem.Reference"
      DisplayName="Log: {step['name']} started" />

    <!-- ════════════════════════════════════════════════════════════ -->
    <!-- IMPLEMENTATION — Pseudo-steps as executable structure        -->
    <!-- ════════════════════════════════════════════════════════════ -->
{pseudo_steps_xml}
{business_rule_xml}
    <!-- ════════════════════════════════════════════════════════════ -->

    <ui:LogMessage Level="Trace"
      Message="'[{process_name}] {step['name']} - Completed | TxRef: ' + in_TransactionItem.Reference"
      DisplayName="Log: {step['name']} completed" />

  </Sequence>
</Activity>
"""


def generate_open_app_xaml(app: str, process_name: str, config_settings: list) -> str:
    """Generates Business/OpenAppName.xaml with credential handling hints and standardized logs."""
    url_key = f"{app}_URL"
    cred_key = f"{app}_CredentialAsset"
    url_val = next((s["value"] for s in config_settings if s["name"] == url_key), "TODO")

    return f"""<?xml version="1.0" encoding="utf-8"?>
<!--
  Open{app}.xaml
  Opens and initializes {app}, ready for transaction processing.
  Generated by ReFramework Generator v{VERSION} on {datetime.now().strftime("%Y-%m-%d")}
-->
<Activity
  mc:Ignorable="sap sap2010"
  x:Class="UiPath.ReFramework.Business.Open{app}"
  xmlns="http://schemas.microsoft.com/netfx/2009/xaml/activities"
  xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
  xmlns:mva="clr-namespace:Microsoft.VisualBasic.Activities;assembly=System.Activities"
  xmlns:sap2010="http://schemas.microsoft.com/netfx/2010/xaml/activities/presentation"
  xmlns:scg="clr-namespace:System.Collections.Generic;assembly=mscorlib"
  xmlns:s="clr-namespace:System;assembly=mscorlib"
  xmlns:ui="http://schemas.uipath.com/workflow/activities"
  xmlns:x="http://schemas.microsoft.com/winfx/2006/xaml">
  <x:Members>
    <x:Property Name="in_Config" Type="InArgument(scg:Dictionary(x:String, x:Object))"
      sap2010:Annotation.AnnotationText="Reads: Config('{url_key}'), Config('{cred_key}') | Output: N/A | Throws: ApplicationException if login fails" />
  </x:Members>
  <Sequence
    sap2010:Annotation.AnnotationText="Open and initialize {app}&#xA;URL: in_Config(&amp;quot;{url_key}&amp;quot;) → default: {url_val}&#xA;Credentials: Get Credentials asset from in_Config(&amp;quot;{cred_key}&amp;quot;)&#xA;&#xA;Reads: Config('{url_key}'), Config('{cred_key}') | Output: N/A | Throws: ApplicationException if login fails"
    DisplayName="Open {app}">

    <ui:LogMessage Level="Info"
      Message="'[{process_name}] Init - Opening {app}'"
      DisplayName="Log: Init - Opening {app}" />

    <!-- ════════════════════════════════════════════════════════════ -->
    <!-- STEP 1: Get credentials from Orchestrator Asset             -->
    <!-- Use "Get Credentials" activity:                             -->
    <!--   Asset Name: in_Config("{cred_key}").ToString()            -->
    <!--   Output: username (String), password (SecureString)        -->
    <!-- ════════════════════════════════════════════════════════════ -->
    <Sequence DisplayName="TODO: Get credentials"
      sap2010:Annotation.AnnotationText="TODO: implement - Get credentials from Orchestrator asset Config('{cred_key}')">
      <!-- TODO: implement - Get credentials from Orchestrator -->
    </Sequence>

    <!-- ════════════════════════════════════════════════════════════ -->
    <!-- STEP 2: Open the application                                -->
    <!-- Option A — Web app: Use "Open Browser" activity             -->
    <!--   URL: in_Config("{url_key}").ToString()                    -->
    <!--   Browser: Chrome (recommended)                             -->
    <!-- Option B — Desktop app: Use "Start Process" or             -->
    <!--   "Open Application" activity                               -->
    <!-- ════════════════════════════════════════════════════════════ -->
    <Sequence DisplayName="TODO: Open application"
      sap2010:Annotation.AnnotationText="TODO: implement - Open {app} using Config('{url_key}')">
      <!-- TODO: implement - Open browser/application -->
    </Sequence>

    <!-- ════════════════════════════════════════════════════════════ -->
    <!-- STEP 3: Login                                               -->
    <!-- Use "Type Into" for username field                          -->
    <!-- Use "Type Into" with SendWindowMessages for password        -->
    <!-- Use "Click" for Login/Submit button                         -->
    <!-- ════════════════════════════════════════════════════════════ -->
    <Sequence DisplayName="TODO: Login"
      sap2010:Annotation.AnnotationText="TODO: implement - Enter credentials and click login button">
      <!-- TODO: implement - Login flow -->
    </Sequence>

    <!-- ════════════════════════════════════════════════════════════ -->
    <!-- STEP 4: Verify successful login                             -->
    <!-- Use "Check App State" or "Element Exists" to confirm        -->
    <!-- dashboard/home page loaded. If not, throw:                  -->
    <!--   New ApplicationException("{app} login failed")            -->
    <!-- ════════════════════════════════════════════════════════════ -->
    <Sequence DisplayName="TODO: Verify login success"
      sap2010:Annotation.AnnotationText="TODO: implement - Verify login succeeded. Throw ApplicationException if failed.">
      <!-- TODO: implement - Verify login, throw ApplicationException if failed -->
    </Sequence>

    <ui:LogMessage Level="Info"
      Message="'[{process_name}] Init - {app} ready'"
      DisplayName="Log: Init - {app} ready" />

  </Sequence>
</Activity>
"""


def generate_close_all_applications_xaml(apps: list, process_name: str) -> str:
    """Generates Framework/CloseAllApplications.xaml with standardized log messages."""
    close_blocks = ""
    for app in apps:
        close_blocks += f"""
    <!-- Close {app} -->
    <TryCatch DisplayName="Try close {app} (ignore errors)">
      <TryCatch.Try>
        <Sequence DisplayName="Close {app}">
          <ui:LogMessage Level="Trace"
            Message="'[{process_name}] Close - Closing {app}'"
            DisplayName="Log: Close - Closing {app}" />
          <!-- TODO: Add Close Tab / Close Browser activity for {app} -->
          <ui:LogMessage Level="Trace"
            Message="'[{process_name}] Close - {app} closed'"
            DisplayName="Log: Close - {app} closed" />
        </Sequence>
      </TryCatch.Try>
      <TryCatch.Catches>
        <Catch x:TypeArguments="s:Exception">
          <ActivityAction x:TypeArguments="s:Exception">
            <ActivityAction.Argument>
              <DelegateInArgument x:TypeArguments="s:Exception" Name="closeEx" />
            </ActivityAction.Argument>
            <ui:LogMessage Level="Warn"
              Message="['[{process_name}] Close - Close failed (non-critical): ' + closeEx.Message]"
              DisplayName="Log: Close failed (non-critical)" />
          </ActivityAction>
        </Catch>
      </TryCatch.Catches>
    </TryCatch>"""

    return f"""<?xml version="1.0" encoding="utf-8"?>
<Activity
  mc:Ignorable="sap sap2010"
  x:Class="UiPath.ReFramework.Framework.CloseAllApplications"
  xmlns="http://schemas.microsoft.com/netfx/2009/xaml/activities"
  xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
  xmlns:sap2010="http://schemas.microsoft.com/netfx/2010/xaml/activities/presentation"
  xmlns:scg="clr-namespace:System.Collections.Generic;assembly=mscorlib"
  xmlns:s="clr-namespace:System;assembly=mscorlib"
  xmlns:ui="http://schemas.uipath.com/workflow/activities"
  xmlns:x="http://schemas.microsoft.com/winfx/2006/xaml">
  <x:Members>
    <x:Property Name="in_Config" Type="InArgument(scg:Dictionary(x:String, x:Object))"
      sap2010:Annotation.AnnotationText="Reads: N/A | Output: N/A | Throws: N/A (all exceptions caught)" />
  </x:Members>
  <Sequence DisplayName="Close All Applications"
    sap2010:Annotation.AnnotationText="Gracefully close all applications used in {process_name}.&#xA;Each close is wrapped in TryCatch to prevent End Process from failing.&#xA;&#xA;Reads: N/A | Output: N/A | Throws: N/A (all exceptions caught)">
    {close_blocks}
  </Sequence>
</Activity>
"""


def generate_kill_all_processes_xaml(apps: list, process_name: str) -> str:
    """Generates Framework/KillAllProcesses.xaml with standardized log messages."""
    kill_blocks = ""
    for app in apps:
        # Generate a reasonable process executable name guess
        app_lower = app.lower()
        if "chrome" in app_lower or "web" in app_lower or "browser" in app_lower:
            process_exe = "chrome.exe"
        elif "excel" in app_lower:
            process_exe = "EXCEL.EXE"
        elif "sap" in app_lower:
            process_exe = "saplogon.exe"
        elif "outlook" in app_lower:
            process_exe = "OUTLOOK.EXE"
        else:
            process_exe = f"{app}.exe"

        kill_blocks += f"""
    <!-- Kill {app} process if still running -->
    <TryCatch DisplayName="Kill {app} if running">
      <TryCatch.Try>
        <Sequence DisplayName="Terminate {app}">
          <ui:LogMessage Level="Trace"
            Message="'[{process_name}] Kill - Terminating {process_exe}'"
            DisplayName="Log: Kill - Terminating {process_exe}" />
          <!-- TODO: Add Kill Process activity -->
          <!-- ProcessName: "{process_exe}" (without .exe if using UiPath activity) -->
        </Sequence>
      </TryCatch.Try>
      <TryCatch.Catches>
        <Catch x:TypeArguments="s:Exception">
          <ActivityAction x:TypeArguments="s:Exception">
            <ActivityAction.Argument>
              <DelegateInArgument x:TypeArguments="s:Exception" Name="killEx" />
            </ActivityAction.Argument>
            <ui:LogMessage Level="Trace"
              Message="['[{process_name}] Kill - {app} process not running: ' + killEx.Message]"
              DisplayName="Log: Process not running" />
          </ActivityAction>
        </Catch>
      </TryCatch.Catches>
    </TryCatch>"""

    return f"""<?xml version="1.0" encoding="utf-8"?>
<Activity
  mc:Ignorable="sap sap2010"
  x:Class="UiPath.ReFramework.Framework.KillAllProcesses"
  xmlns="http://schemas.microsoft.com/netfx/2009/xaml/activities"
  xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
  xmlns:sap2010="http://schemas.microsoft.com/netfx/2010/xaml/activities/presentation"
  xmlns:s="clr-namespace:System;assembly=mscorlib"
  xmlns:ui="http://schemas.uipath.com/workflow/activities"
  xmlns:x="http://schemas.microsoft.com/winfx/2006/xaml">
  <Sequence DisplayName="Kill All Processes"
    sap2010:Annotation.AnnotationText="Force-kill all application processes.&#xA;Called at Init start (clean state) and End Process.&#xA;Each kill is wrapped in TryCatch.&#xA;&#xA;Reads: N/A | Output: N/A | Throws: N/A (all exceptions caught)">
    {kill_blocks}
  </Sequence>
</Activity>
"""


def generate_set_transaction_status_xaml(process_name: str) -> str:
    """
    Generates Framework/SetTransactionStatus.xaml with three explicit branches:
    - Success: marks Successful, logs completion
    - BusinessRuleException: marks Failed, logs business rule violation, NO retry (does not rethrow)
    - ApplicationException: rethrows so Main.xaml can handle retry logic
    """
    return f"""<?xml version="1.0" encoding="utf-8"?>
<Activity
  mc:Ignorable="sap sap2010"
  x:Class="UiPath.ReFramework.Framework.SetTransactionStatus"
  xmlns="http://schemas.microsoft.com/netfx/2009/xaml/activities"
  xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
  xmlns:mva="clr-namespace:Microsoft.VisualBasic.Activities;assembly=System.Activities"
  xmlns:sap2010="http://schemas.microsoft.com/netfx/2010/xaml/activities/presentation"
  xmlns:scg="clr-namespace:System.Collections.Generic;assembly=mscorlib"
  xmlns:ui="http://schemas.uipath.com/workflow/activities"
  xmlns:uia="clr-namespace:UiPath.Core;assembly=UiPath.Core"
  xmlns:s="clr-namespace:System;assembly=mscorlib"
  xmlns:x="http://schemas.microsoft.com/winfx/2006/xaml">
  <x:Members>
    <x:Property Name="in_TransactionItem" Type="InArgument(ui:QueueItem)"
      sap2010:Annotation.AnnotationText="The queue item whose status will be set." />
    <x:Property Name="in_TransactionStatus" Type="InArgument(x:String)"
      sap2010:Annotation.AnnotationText="Status to set: 'Successful', 'BusinessException', or 'SystemException'." />
    <x:Property Name="in_TransactionError" Type="InArgument(s:Exception)"
      sap2010:Annotation.AnnotationText="Exception object (required for BusinessException and SystemException branches)." />
    <x:Property Name="in_Config" Type="InArgument(scg:Dictionary(x:String, x:Object))"
      sap2010:Annotation.AnnotationText="Reads: N/A | Output: N/A | Throws: N/A" />
  </x:Members>
  <Sequence DisplayName="Set Transaction Status"
    sap2010:Annotation.AnnotationText="Sets Orchestrator queue item status.&#xA;&#xA;Branch Success: Status=Successful, logs completion&#xA;Branch BusinessRuleException: Status=Failed, NO retry, does NOT rethrow&#xA;Branch ApplicationException: Rethrows to Main.xaml for retry handling">

    <ui:LogMessage Level="Trace"
      Message="['[{process_name}] SetStatus - ' + in_TransactionItem.Reference + ' → ' + in_TransactionStatus + ' | ' + If(in_TransactionError IsNot Nothing, in_TransactionError.Message, &quot;OK&quot;)]"
      DisplayName="Log: SetStatus entry" />

    <Switch x:TypeArguments="x:String" Expression="[in_TransactionStatus]" DisplayName="Switch on transaction status">
      <Switch.Cases>

        <!-- ══════════════════════════════════════════════════════════════ -->
        <!-- BRANCH: SUCCESS                                                -->
        <!-- Sets queue item to Successful, logs completion message         -->
        <!-- ══════════════════════════════════════════════════════════════ -->
        <Case x:Key="Successful">
          <Sequence DisplayName="Branch: Success"
            sap2010:Annotation.AnnotationText="Reads: N/A | Output: N/A | Throws: N/A">
            <ui:SetTransactionStatus DisplayName="Set Queue Item → Successful"
              QueueItem="[in_TransactionItem]"
              Status="Successful" />
            <ui:LogMessage Level="Info"
              Message="['[{process_name}] Transaction ' + in_TransactionItem.Reference + ' completed successfully']"
              DisplayName="Log: Transaction completed successfully" />
          </Sequence>
        </Case>

        <!-- ══════════════════════════════════════════════════════════════ -->
        <!-- BRANCH: BUSINESS RULE EXCEPTION                                -->
        <!-- Sets queue item to Failed with reason, NO retry triggered      -->
        <!-- Does NOT rethrow — exception is fully handled here             -->
        <!-- ══════════════════════════════════════════════════════════════ -->
        <Case x:Key="BusinessException">
          <Sequence DisplayName="Branch: BusinessRuleException (no retry)"
            sap2010:Annotation.AnnotationText="Reads: N/A | Output: N/A | Throws: N/A (exception is consumed, not rethrown)">
            <ui:SetTransactionStatus DisplayName="Set Queue Item → Failed (Business)"
              QueueItem="[in_TransactionItem]"
              Status="Failed"
              Reason="[in_TransactionError.Message]"
              ErrorType="Business" />
            <ui:LogMessage Level="Warn"
              Message="['[{process_name}] Business rule violation on ' + in_TransactionItem.Reference + ': ' + in_TransactionError.Message]"
              DisplayName="Log: Business rule violation" />
            <!-- NO Rethrow here — BusinessRuleException does not trigger retry -->
          </Sequence>
        </Case>

        <!-- ══════════════════════════════════════════════════════════════ -->
        <!-- BRANCH: APPLICATION / SYSTEM EXCEPTION                         -->
        <!-- Rethrows the exception so Main.xaml can handle retry logic     -->
        <!-- Main.xaml increments retry counter and may go back to Init     -->
        <!-- ══════════════════════════════════════════════════════════════ -->
        <Case x:Key="SystemException">
          <Sequence DisplayName="Branch: ApplicationException (rethrow for retry)"
            sap2010:Annotation.AnnotationText="Reads: N/A | Output: N/A | Throws: ApplicationException (rethrown to Main.xaml)">
            <ui:SetTransactionStatus DisplayName="Set Queue Item → Failed (Application)"
              QueueItem="[in_TransactionItem]"
              Status="Failed"
              Reason="[in_TransactionError.Message]"
              ErrorType="Application" />
            <ui:LogMessage Level="Error"
              Message="['[{process_name}] Application exception on ' + in_TransactionItem.Reference + ': ' + in_TransactionError.Message + ' — rethrowing for retry']"
              DisplayName="Log: Application exception (will retry)" />
            <!-- Rethrow so Main.xaml can handle retry logic -->
            <Rethrow DisplayName="Rethrow to Main.xaml for retry handling" />
          </Sequence>
        </Case>

      </Switch.Cases>
    </Switch>

  </Sequence>
</Activity>
"""


def generate_init_all_settings_xaml(process_name: str) -> str:
    """Generates Framework/InitAllSettings.xaml with standardized log messages."""
    return f"""<?xml version="1.0" encoding="utf-8"?>
<Activity
  mc:Ignorable="sap sap2010"
  x:Class="UiPath.ReFramework.Framework.InitAllSettings"
  xmlns="http://schemas.microsoft.com/netfx/2009/xaml/activities"
  xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
  xmlns:mva="clr-namespace:Microsoft.VisualBasic.Activities;assembly=System.Activities"
  xmlns:sap2010="http://schemas.microsoft.com/netfx/2010/xaml/activities/presentation"
  xmlns:scg="clr-namespace:System.Collections.Generic;assembly=mscorlib"
  xmlns:s="clr-namespace:System;assembly=mscorlib"
  xmlns:ui="http://schemas.uipath.com/workflow/activities"
  xmlns:x="http://schemas.microsoft.com/winfx/2006/xaml">
  <x:Members>
    <x:Property Name="in_ConfigFile" Type="InArgument(x:String)"
      sap2010:Annotation.AnnotationText="Relative path to Config.xlsx. Default: Data\\Config.xlsx | Reads: N/A | Output: N/A | Throws: N/A" />
    <x:Property Name="in_ConfigSheets" Type="InArgument(s:String[])"
      sap2010:Annotation.AnnotationText="Sheet names to read. Default: [&amp;quot;Settings&amp;quot;, &amp;quot;Constants&amp;quot;] | Reads: N/A | Output: N/A | Throws: N/A" />
    <x:Property Name="out_Config" Type="OutArgument(scg:Dictionary(x:String, x:Object))"
      sap2010:Annotation.AnnotationText="Output: Dictionary populated with all settings and constants. | Reads: Config.xlsx | Output: out_Config (Dictionary) | Throws: ApplicationException if file not found" />
  </x:Members>
  <Sequence DisplayName="Init All Settings"
    sap2010:Annotation.AnnotationText="Reads Config.xlsx and returns a Dictionary(String, Object).&#xA;Also reads Orchestrator Assets if connected.&#xA;Called ONCE at process start (first Init state only).&#xA;&#xA;Reads: Data/Config.xlsx | Output: out_Config (Dictionary) | Throws: ApplicationException if config load fails">

    <ui:LogMessage Level="Info"
      Message="'[{process_name}] Init - Loading configuration'"
      DisplayName="Log: Init - Loading configuration" />

    <TryCatch DisplayName="Try load configuration">
      <TryCatch.Try>
        <Sequence DisplayName="Load Config.xlsx">
          <!-- ── Read Settings sheet ── -->
          <ui:ExcelApplicationScope DisplayName="Open Config.xlsx"
            WorkbookPath="[in_ConfigFile]">
            <ui:ExcelApplicationScope.Body>
              <ActivityAction>
                <Sequence DisplayName="Read sheets">
                  <!-- TODO: Use ReadRange activity for each sheet in in_ConfigSheets -->
                  <!-- Then iterate rows and populate out_Config dictionary:           -->
                  <!-- out_Config(row("Name").ToString()) = row("Value")               -->
                </Sequence>
              </ActivityAction>
            </ui:ExcelApplicationScope.Body>
          </ui:ExcelApplicationScope>

          <!-- ── Override with Orchestrator Assets (if connected) ── -->
          <!-- TODO (optional): Use GetAsset activity for sensitive values -->
          <!-- This overrides Config.xlsx values with Orchestrator-managed assets -->

          <ui:LogMessage Level="Info"
            Message="['[{process_name}] Init - Configuration loaded. Keys: ' + out_Config.Count.ToString()]"
            DisplayName="Log: Init - Configuration loaded" />
        </Sequence>
      </TryCatch.Try>
      <TryCatch.Catches>
        <Catch x:TypeArguments="s:Exception">
          <ActivityAction x:TypeArguments="s:Exception">
            <ActivityAction.Argument>
              <DelegateInArgument x:TypeArguments="s:Exception" Name="configException" />
            </ActivityAction.Argument>
            <Sequence DisplayName="Handle config load error">
              <ui:LogMessage Level="Error"
                Message="['[{process_name}] Init - Failed to load config: ' + configException.Message]"
                DisplayName="Log: Init - Failed to load config" />
              <Throw Exception="[New ApplicationException(&quot;[{process_name}] Failed to load configuration: &quot; + configException.Message, configException)]"
                DisplayName="Throw ApplicationException" />
            </Sequence>
          </ActivityAction>
        </Catch>
      </TryCatch.Catches>
    </TryCatch>

  </Sequence>
</Activity>
"""


# ─── STEP 5: GENERATE CONFIG.XLSX ─────────────────────────────────────────────

def generate_config_xlsx(metadata: dict, output_path: str):
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        print("  ⚠️  openpyxl not installed - Config.xlsx skipped (install with: pip install openpyxl)")
        return

    wb = openpyxl.Workbook()
    process_name = metadata["process_name"]

    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill("solid", fgColor="1F4E79")
    alt_fill = PatternFill("solid", fgColor="D6E4F0")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin")
    )

    def write_sheet(ws, rows, title_note=""):
        headers = ["Name", "Value", "Description"]
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_border
        ws.row_dimensions[1].height = 20

        for row_idx, (name, value, desc) in enumerate(rows, 2):
            fill = alt_fill if row_idx % 2 == 0 else None
            for col, val in enumerate([name, value, desc], 1):
                cell = ws.cell(row=row_idx, column=col, value=val)
                cell.border = thin_border
                if fill:
                    cell.fill = fill

        ws.column_dimensions["A"].width = 38
        ws.column_dimensions["B"].width = 32
        ws.column_dimensions["C"].width = 50
        ws.freeze_panes = "A2"

    # ── Settings sheet ──
    ws_settings = wb.active
    ws_settings.title = "Settings"
    settings_rows = [
        ("OrchestratorQueueName", metadata.get("queue_name", ""), "Orchestrator queue name (overrides --input)"),
        ("MaxRetryNumber", metadata.get("max_retry_number", 3), "Max retries before System Exception aborts transaction"),
        ("MaxConsecutiveSystemExceptions", 3, "Max consecutive system exceptions before aborting entire process"),
        ("logF_BusinessProcessName", process_name, "Process name tag added to all Orchestrator logs"),
    ]
    for app in metadata.get("applications", []):
        settings_rows.append((f"{app}_URL", "TODO: add URL", f"Base URL / entry point for {app}"))
        settings_rows.append((f"{app}_CredentialAsset", f"{app}_Credentials", f"Orchestrator Asset name for {app} credentials"))

    write_sheet(ws_settings, settings_rows)

    # ── Constants sheet ──
    ws_constants = wb.create_sheet("Constants")
    constants_rows = []
    for s in metadata.get("config_settings", []):
        if s["type"] in ("Constant", "Asset"):
            constants_rows.append((s["name"], s["value"], s.get("description", f"Type: {s['type']}")))

    # Add transaction item fields as reference constants
    for f in metadata.get("transaction_item_fields", []):
        if isinstance(f, dict):
            constants_rows.append((
                f"TxField_{f['name']}",
                f['name'],
                f"Transaction item field — {f.get('description', '')} (Type: {f.get('type', 'String')})"
            ))

    if not constants_rows:
        constants_rows = [("EXAMPLE_CONSTANT", "replace_me", "Add business rule constants and thresholds here")]

    write_sheet(ws_constants, constants_rows)

    wb.save(output_path)
    print(f"  ✓ Data/Config.xlsx ({len(settings_rows)} settings, {len(constants_rows)} constants)")


# ─── STEP 6: ASSEMBLE FULL PROJECT ───────────────────────────────────────────

def build_project(metadata: dict, output_base: str, variant: str = "reframework") -> str:
    process_name = metadata["process_name"]
    project_dir = Path(output_base) / process_name
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "Framework").mkdir(exist_ok=True)
    (project_dir / "Business").mkdir(exist_ok=True)
    (project_dir / "Data").mkdir(exist_ok=True)

    print(f"\n📁 Assembling project: {process_name}")
    print(f"   Variant: {variant}  |  Steps: {len(metadata.get('process_steps', []))}  |  "
          f"Apps: {len(metadata.get('applications', []))}  |  "
          f"Source: {metadata.get('transaction_source', 'Queue')}")

    apps = metadata.get("applications", [])
    config_settings = metadata.get("config_settings", [])

    # ── project.json ──
    project_json = {
        "name": process_name,
        "description": metadata["process_description"],
        "main": "Main.xaml",
        "dependencies": {
            "UiPath.Excel.Activities": "[2.22.0, )",
            "UiPath.Mail.Activities": "[1.20.0, )",
            "UiPath.System.Activities": "[24.10.0, )",
            "UiPath.UIAutomation.Activities": "[24.10.0, )"
        },
        "webServiceReferences": {},
        "schemaVersion": "4.0",
        "studioVersion": "24.10.0.0",
        "projectVersion": "1.0.0",
        "runtimeOptions": {"autoDispose": False, "isPausable": True},
        "designOptions": {
            "projectProfile": "Unattended",
            "outputType": "Process",
            "libraryOptions": {"includeOriginalXaml": False},
            "workflowAnalyzerConfig": {}
        }
    }
    (project_dir / "project.json").write_text(json.dumps(project_json, indent=2), encoding="utf-8")
    print("  ✓ project.json")

    # ── Main.xaml ──
    if variant == "sequence":
        # Use sequence template instead of state machine
        seq_template = SEQUENCE_TEMPLATES_DIR / "Main.xaml"
        if seq_template.exists():
            main_content = seq_template.read_text(encoding="utf-8")
        else:
            main_content = load_template("Main.xaml")  # Fallback
        main_content = apply_common_replacements(main_content, metadata)
        # Add step invocations for sequence
        steps_xml = "\n".join(
            generate_process_step_invocation(step, process_name)
            for step in metadata.get("process_steps", [])
        ) or "          <!-- TODO: Add your process steps here -->"
        main_content = main_content.replace("{{PROCESS_STEPS_PLACEHOLDER}}", steps_xml)
        (project_dir / "Main.xaml").write_text(main_content, encoding="utf-8")
        print("  ✓ Main.xaml (Sequence)")
    else:
        main_content = load_template("Main.xaml")
        main_content = apply_common_replacements(main_content, metadata)
        (project_dir / "Main.xaml").write_text(main_content, encoding="utf-8")
        print("  ✓ Main.xaml (State Machine)")

    # ── Process.xaml ──
    steps_xml = "\n".join(
        generate_process_step_invocation(step, process_name)
        for step in metadata.get("process_steps", [])
    ) or "    <!-- TODO: Add your business step invocations here -->"

    process_content = load_template("Process.xaml")
    process_content = apply_common_replacements(process_content, metadata)
    process_content = process_content.replace("{{PROCESS_STEPS_PLACEHOLDER}}", steps_xml)
    (project_dir / "Process.xaml").write_text(process_content, encoding="utf-8")
    print(f"  ✓ Process.xaml ({len(metadata.get('process_steps', []))} step invocations)")

    # ── Framework/InitAllApplications.xaml ──
    app_stubs = "\n".join(generate_app_init_invocation(app) for app in apps) \
                or "    <!-- TODO: Add application init steps -->"
    init_apps = load_template("InitAllApplications.xaml")
    init_apps = apply_common_replacements(init_apps, metadata)
    init_apps = init_apps.replace("{{APPLICATION_STUBS_PLACEHOLDER}}", app_stubs)
    (project_dir / "Framework" / "InitAllApplications.xaml").write_text(init_apps, encoding="utf-8")
    print("  ✓ Framework/InitAllApplications.xaml")

    # ── Framework/GetTransactionData.xaml ──
    get_tx_body = generate_get_transaction_body(metadata, process_name)
    get_tx = load_template("GetTransactionData.xaml")
    get_tx = apply_common_replacements(get_tx, metadata)
    get_tx = get_tx.replace("{{GET_TRANSACTION_BODY}}", get_tx_body)
    (project_dir / "Framework" / "GetTransactionData.xaml").write_text(get_tx, encoding="utf-8")
    print(f"  ✓ Framework/GetTransactionData.xaml ({metadata.get('transaction_source', 'Queue')} mode)")

    # ── Framework/SetTransactionStatus.xaml ──
    (project_dir / "Framework" / "SetTransactionStatus.xaml").write_text(
        generate_set_transaction_status_xaml(process_name), encoding="utf-8")
    print("  ✓ Framework/SetTransactionStatus.xaml")

    # ── Framework/InitAllSettings.xaml ──
    (project_dir / "Framework" / "InitAllSettings.xaml").write_text(
        generate_init_all_settings_xaml(process_name), encoding="utf-8")
    print("  ✓ Framework/InitAllSettings.xaml")

    # ── Framework/CloseAllApplications.xaml ──
    (project_dir / "Framework" / "CloseAllApplications.xaml").write_text(
        generate_close_all_applications_xaml(apps, process_name), encoding="utf-8")
    print("  ✓ Framework/CloseAllApplications.xaml")

    # ── Framework/KillAllProcesses.xaml ──
    (project_dir / "Framework" / "KillAllProcesses.xaml").write_text(
        generate_kill_all_processes_xaml(apps, process_name), encoding="utf-8")
    print("  ✓ Framework/KillAllProcesses.xaml")

    # ── Business step stubs ──
    for step in metadata.get("process_steps", []):
        content = generate_business_step_xaml(step, process_name)
        (project_dir / "Business" / f"{step['name']}.xaml").write_text(content, encoding="utf-8")
        print(f"  ✓ Business/{step['name']}.xaml  [{len(step.get('pseudo_steps', []))} pseudo-steps]")

    # ── Business open-app stubs ──
    for app in apps:
        content = generate_open_app_xaml(app, process_name, config_settings)
        (project_dir / "Business" / f"Open{app}.xaml").write_text(content, encoding="utf-8")
        print(f"  ✓ Business/Open{app}.xaml")

    # ── Data/Config.xlsx ──
    generate_config_xlsx(metadata, str(project_dir / "Data" / "Config.xlsx"))

    # ── extracted_metadata.json (for reference / debugging) ──
    (project_dir / "Data" / "extracted_metadata.json").write_text(
        json.dumps(metadata, indent=2, ensure_ascii=False), encoding="utf-8")
    print("  ✓ Data/extracted_metadata.json")

    # ── README.md ──
    (project_dir / "README.md").write_text(generate_project_readme(metadata), encoding="utf-8")
    print("  ✓ README.md")

    return str(project_dir)


def generate_project_readme(metadata: dict) -> str:
    steps_md = "\n".join(
        f"| `{s['id']}` | **{s['name']}** | {s['app']} | {s['description']} |"
        for s in metadata.get("process_steps", [])
    )
    apps_md = ", ".join(f"`{a}`" for a in metadata.get("applications", []))
    biz_exc_md = "\n".join(
        f"- `{e['name']}`: {e['condition']}"
        for e in metadata.get("business_exceptions", [])
    ) or "_None identified_"
    sys_exc_md = "\n".join(
        f"- `{e['name']}`: {e['condition']} → _{e.get('recovery_hint', 'Retry via ReFramework')}_"
        for e in metadata.get("system_exceptions", [])
    ) or "_None identified_"

    return f"""# {metadata['process_name']}

> {metadata['process_description']}

**Generated** by [ReFramework Generator](https://github.com/alessiodonato/uipath-reframework-generator) on {datetime.now().strftime("%Y-%m-%d")}
**Transaction Source:** `{metadata.get('transaction_source', 'Queue')}` | **Max Retries:** `{metadata.get('max_retry_number', 3)}`

---

## Applications
{apps_md}

## Process Steps

| ID | Step | Application | Description |
|----|------|-------------|-------------|
{steps_md}

## Exception Handling

### Business Exceptions (no retry)
{biz_exc_md}

### System Exceptions (triggers retry)
{sys_exc_md}

---

## Getting Started

1. **Open in UiPath Studio**: `File → Open Project → select project.json`
2. **Install NuGet packages**: Studio will prompt automatically on first open
3. **Configure** `Data/Config.xlsx` — fill in URLs and credential asset names
4. **Implement** `Business/*.xaml` — search for `<!-- TODO:` comments
5. **Test** with a single queue item before deploying to Orchestrator

## Implementation Checklist

- [ ] `Data/Config.xlsx` — fill all `TODO: add URL` values
- [ ] `Data/Config.xlsx` — verify credential asset names match Orchestrator
{"".join(f"- [ ] `Business/Open{a}.xaml` — add Open Browser / login selectors{chr(10)}" for a in metadata.get("applications", []))}{"".join(f"- [ ] `Business/{s['name']}.xaml` — implement {len(s.get('pseudo_steps', []))} steps{chr(10)}" for s in metadata.get("process_steps", []))}
## Files Generated

```
{metadata['process_name']}/
├── project.json                          # Open this in UiPath Studio
├── Main.xaml                             # ✅ Complete State Machine
├── Process.xaml                          # ✅ Step invocations wired
├── Data/
│   ├── Config.xlsx                       # ⚠️  Fill in URLs and credentials
│   └── extracted_metadata.json          # Reference: AI extraction output
├── Framework/
│   ├── InitAllSettings.xaml             # ⚠️  Review TODO comments
│   ├── InitAllApplications.xaml         # ✅ App stubs wired
│   ├── GetTransactionData.xaml          # ✅ {metadata.get('transaction_source', 'Queue')} mode configured
│   ├── SetTransactionStatus.xaml        # ✅ Complete
│   ├── CloseAllApplications.xaml        # ⚠️  Add Close Browser/App selectors
│   └── KillAllProcesses.xaml            # ⚠️  Add process names to kill
└── Business/
{"".join(f"    ├── Open{a}.xaml{chr(10)}" for a in metadata.get("applications", []))}{"".join(f"    ├── {s['name']}.xaml{chr(10)}" for s in metadata.get("process_steps", []))}```

Legend: ✅ Ready to use | ⚠️ Needs configuration | 🔧 Needs implementation
"""


# ─── STEP 7: ZIP AND DELIVER ──────────────────────────────────────────────────

def zip_project(project_dir: str, output_dir: str) -> str:
    process_name = Path(project_dir).name
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    zip_path = str(Path(output_dir) / f"{process_name}_ReFramework.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for file_path in Path(project_dir).rglob("*"):
            if file_path.is_file():
                zf.write(file_path, file_path.relative_to(Path(project_dir).parent))
    size_kb = Path(zip_path).stat().st_size // 1024
    print(f"\n📦 ZIP created: {zip_path} ({size_kb} KB)")
    return zip_path


# ─── INTERACTIVE MODE ─────────────────────────────────────────────────────────

def interactive_mode() -> argparse.Namespace:
    print("╔══════════════════════════════════════════════════════╗")
    print("║    UiPath ReFramework Generator — Interactive Mode   ║")
    print("╚══════════════════════════════════════════════════════╝\n")

    input_path = input("📄 Path to your PDD (PDF or DOCX): ").strip().strip('"')
    if not input_path:
        print("Error: no input file provided.")
        sys.exit(1)

    output_dir = input("📁 Output directory [./output]: ").strip() or "./output"
    api_key = input("🔑 Anthropic API key [ANTHROPIC_API_KEY env var]: ").strip() or ""

    args = argparse.Namespace(
        input=input_path,
        output=output_dir,
        api_key=api_key or None,
        metadata_only=False
    )
    return args


# ─── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Generate a complete UiPath ReFramework project from a PDD (PDF or DOCX)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python generate_reframework.py                                    # interactive
  python generate_reframework.py --input Invoice_PDD.pdf
  python generate_reframework.py --input Process_PDD.docx --output ./projects/
  python generate_reframework.py --input Process_PDD.pdf --metadata-only
        """
    )
    parser.add_argument("--input", help="Path to PDD file (PDF or DOCX)")
    parser.add_argument("--output", default="./output", help="Output directory (default: ./output)")
    parser.add_argument("--api-key", help="Anthropic API key (or set ANTHROPIC_API_KEY env var)")
    parser.add_argument("--metadata-only", action="store_true",
                        help="Only print extracted metadata JSON, don't generate files")
    parser.add_argument("--variant", choices=list(VARIANTS.keys()), default="reframework",
                        help="Project variant (default: reframework)")
    parser.add_argument("--validate", action="store_true",
                        help="Validate generated XAML files after generation")
    parser.add_argument("--resolve-nuget", action="store_true",
                        help="Query NuGet API for latest package versions")
    parser.add_argument("--version", action="version", version=f"ReFramework Generator v{VERSION}")

    args = parser.parse_args()

    # Interactive mode if no input provided
    if not args.input:
        args = interactive_mode()

    # Resolve API key: --api-key flag > ANTHROPIC_API_KEY env var
    api_key = args.api_key or os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        print("\n❌ ERROR: No Anthropic API key found.")
        print("   Set it via: export ANTHROPIC_API_KEY=sk-ant-...")
        print("   Or pass it: --api-key sk-ant-...")
        sys.exit(1)

    # Auto-install dependencies
    ensure_dependencies()

    # Run pipeline
    print(f"\n📄 Reading: {args.input}")
    text = extract_document_text(args.input)
    print(f"   ✓ {len(text):,} characters extracted")

    print("\n🤖 Analyzing document with Claude AI...")
    metadata = extract_metadata_via_api(text, api_key)

    print(f"\n✅ Extraction complete:")
    print(f"   Process  : {metadata['process_name']}")
    print(f"   Steps    : {len(metadata.get('process_steps', []))}")
    print(f"   Apps     : {', '.join(metadata.get('applications', []))}")
    print(f"   Source   : {metadata.get('transaction_source', 'Queue')}")
    print(f"   Biz Exc  : {len(metadata.get('business_exceptions', []))}")
    print(f"   Sys Exc  : {len(metadata.get('system_exceptions', []))}")

    if args.metadata_only:
        print("\n" + json.dumps(metadata, indent=2, ensure_ascii=False))
        return

    # Resolve NuGet versions if requested
    if args.resolve_nuget:
        print("\n📦 Resolving NuGet package versions...")
        try:
            from resolve_nuget import get_latest_versions, COMMON_PACKAGES
            versions = get_latest_versions(COMMON_PACKAGES[:4])
            print(f"   ✓ Resolved {len(versions)} packages")
        except Exception as e:
            print(f"   ⚠️ Could not resolve NuGet versions: {e}")

    # Build project with selected variant
    print(f"\n🏗️ Building project (variant: {args.variant})...")
    project_dir = build_project(metadata, "/tmp/reframework_gen", variant=args.variant)
    zip_path = zip_project(project_dir, args.output)

    total_xaml = len(list(Path(project_dir).rglob("*.xaml")))
    print(f"\n🎉 Done! Generated {total_xaml} XAML files")

    # Validate if requested
    if args.validate:
        print("\n🔍 Validating generated XAML...")
        try:
            from validate_xaml import validate_project
            result = validate_project(project_dir, lint=True)
            print(f"   Files checked: {result.files_checked}")
            print(f"   Errors: {result.error_count} | Warnings: {result.warn_count}")
            if result.has_errors:
                print("   ⚠️ Validation found errors - review before opening in Studio")
            else:
                print("   ✓ All validation checks passed")
        except ImportError:
            print("   ⚠️ Validation module not available")

    print(f"\n   ZIP: {zip_path}")
    print(f"\n📋 Next steps:")
    print(f"   1. Extract the ZIP")
    print(f"   2. Open UiPath Studio → File → Open Project → project.json")
    print(f"   3. Studio will install NuGet packages automatically")
    print(f"   4. Open Data/Config.xlsx and fill in URLs + credential asset names")
    print(f"   5. Search '<!-- TODO:' in XAML files for implementation points")


if __name__ == "__main__":
    main()
